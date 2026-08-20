#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DokuZen Pro - Format Converter
==================================
Konvertiert zwischen verschiedenen Dateiformaten.
"""

from pathlib import Path
from typing import Optional, List, Tuple
from dataclasses import dataclass
from enum import Enum

from utils.logger import LoggerMixin

# Optional imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class OutputFormat(Enum):
    """Ausgabeformate."""
    PDF = "pdf"
    TXT = "txt"
    DOCX = "docx"
    HTML = "html"
    MD = "md"


@dataclass
class ConversionResult:
    """Ergebnis einer Konvertierung."""
    success: bool
    output_path: str
    input_format: str
    output_format: str
    error: Optional[str] = None


class FormatConverter(LoggerMixin):
    """
    Konvertiert Dateien zwischen verschiedenen Formaten.
    
    Unterstützte Konvertierungen:
    - PDF -> TXT, HTML, Images
    - DOCX -> PDF, TXT
    - TXT -> PDF, DOCX
    - Images -> PDF
    """
    
    # Unterstützte Eingabeformate
    SUPPORTED_INPUTS = {".pdf", ".docx", ".doc", ".txt", ".md", ".html",
                        ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"}
    
    def __init__(self):
        self.logger.debug("FormatConverter initialisiert")
    
    def convert(self, input_path: str, output_path: str,
                output_format: Optional[OutputFormat] = None) -> ConversionResult:
        """
        Konvertiert eine Datei.
        
        Args:
            input_path: Eingabedatei
            output_path: Ausgabedatei (oder Verzeichnis)
            output_format: Zielformat (oder aus output_path ermitteln)
            
        Returns:
            ConversionResult
        """
        input_file = Path(input_path)
        output_file = Path(output_path)
        
        if not input_file.exists():
            return ConversionResult(False, output_path, "", "", "Datei nicht gefunden")
        
        input_ext = input_file.suffix.lower()
        
        # Output-Format bestimmen
        if output_format:
            out_ext = f".{output_format.value}"
        else:
            out_ext = output_file.suffix.lower()
        
        if not out_ext:
            return ConversionResult(False, output_path, input_ext, "", 
                                   "Kein Ausgabeformat angegeben")
        
        # Passende Konvertierungsmethode wählen
        if input_ext == ".pdf":
            return self._convert_from_pdf(input_path, output_path, out_ext)
        elif input_ext == ".docx":
            return self._convert_from_docx(input_path, output_path, out_ext)
        elif input_ext in {".txt", ".md"}:
            return self._convert_from_text(input_path, output_path, out_ext)
        elif input_ext == ".html":
            return self._convert_from_html(input_path, output_path, out_ext)
        elif input_ext in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"}:
            return self._convert_from_image(input_path, output_path, out_ext)
        else:
            return ConversionResult(False, output_path, input_ext, out_ext,
                                   f"Format nicht unterstützt: {input_ext}")
    
    def _convert_from_pdf(self, input_path: str, output_path: str,
                          output_ext: str) -> ConversionResult:
        """Konvertiert von PDF."""
        if not PYMUPDF_AVAILABLE:
            return ConversionResult(False, output_path, ".pdf", output_ext,
                                   "PyMuPDF nicht verfügbar")
        
        try:
            doc = fitz.open(input_path)
            try:
                if len(doc) == 0:
                    return ConversionResult(False, output_path, ".pdf", output_ext,
                                           "PDF enthält keine Seiten")

                if output_ext == ".txt":
                    # PDF zu Text
                    text = "\n\n".join(page.get_text() for page in doc)
                    with open(output_path, "w", encoding="utf-8") as f:
                        f.write(text)
                    return ConversionResult(True, output_path, ".pdf", ".txt")

                elif output_ext == ".md":
                    # PDF zu Markdown
                    lines = []
                    for i, page in enumerate(doc):
                        lines.append(f"## Seite {i+1}\n\n{page.get_text()}")
                    text = "\n\n---\n\n".join(lines)
                    with open(output_path, "w", encoding="utf-8") as f:
                        f.write(text)
                    return ConversionResult(True, output_path, ".pdf", ".md")

                elif output_ext == ".html":
                    # PDF zu HTML
                    html_parts = ['<html><head><meta charset="utf-8"></head><body>']
                    for i, page in enumerate(doc):
                        text = page.get_text("html")
                        html_parts.append(f'<div class="page" id="page-{i+1}">{text}</div>')
                    html_parts.append('</body></html>')

                    with open(output_path, "w", encoding="utf-8") as f:
                        f.write("\n".join(html_parts))
                    return ConversionResult(True, output_path, ".pdf", ".html")

                elif output_ext in {".png", ".jpg", ".jpeg", ".bmp", ".webp"}:
                    # PDF zu Bild (erste Seite)
                    page = doc[0]
                    mat = fitz.Matrix(2.0, 2.0)  # 2x Zoom
                    pix = page.get_pixmap(matrix=mat)

                    if output_ext == ".png":
                        pix.save(output_path)
                    else:
                        # Für JPEG/BMP/WEBP via PIL
                        if PIL_AVAILABLE:
                            if pix.n == 1:
                                img = Image.frombytes("L", [pix.width, pix.height], pix.samples).convert("RGB")
                            elif pix.alpha or pix.n == 4:
                                img = Image.frombytes("RGBA", [pix.width, pix.height], pix.samples)
                                bg = Image.new("RGB", img.size, (255, 255, 255))
                                bg.paste(img, mask=img.split()[3])
                                img = bg
                            else:
                                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                            if output_ext in {".jpg", ".jpeg"}:
                                img.save(output_path, "JPEG", quality=90)
                            elif output_ext == ".bmp":
                                img.save(output_path, "BMP")
                            elif output_ext == ".webp":
                                img.save(output_path, "WEBP", quality=90)
                            else:
                                img.save(output_path)
                        else:
                            pix.save(str(Path(output_path).with_suffix(".png")))

                    return ConversionResult(True, output_path, ".pdf", output_ext)

                else:
                    return ConversionResult(False, output_path, ".pdf", output_ext,
                                           f"Konvertierung nicht unterstützt: PDF -> {output_ext}")
            finally:
                doc.close()

        except Exception as e:
            return ConversionResult(False, output_path, ".pdf", output_ext, str(e))
    
    def _convert_from_docx(self, input_path: str, output_path: str,
                           output_ext: str) -> ConversionResult:
        """Konvertiert von DOCX."""
        if not DOCX_AVAILABLE:
            return ConversionResult(False, output_path, ".docx", output_ext,
                                   "python-docx nicht verfügbar")
        
        try:
            doc = DocxDocument(input_path)
            
            if output_ext == ".txt":
                # DOCX zu Text
                text = "\n\n".join(para.text for para in doc.paragraphs)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
                return ConversionResult(True, output_path, ".docx", ".txt")

            elif output_ext == ".md":
                # DOCX zu Markdown
                text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
                return ConversionResult(True, output_path, ".docx", ".md")

            elif output_ext == ".html":
                # DOCX zu HTML
                paras = "".join(f"<p>{para.text}</p>" for para in doc.paragraphs if para.text.strip())
                html = f'<html><head><meta charset="utf-8"></head><body>{paras}</body></html>'
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html)
                return ConversionResult(True, output_path, ".docx", ".html")
            
            elif output_ext == ".pdf":
                # DOCX zu PDF (vereinfacht - nur Text)
                if not REPORTLAB_AVAILABLE:
                    return ConversionResult(False, output_path, ".docx", ".pdf",
                                           "reportlab nicht verfügbar")
                
                c = canvas.Canvas(output_path, pagesize=A4)
                width, height = A4
                y = height - 2*cm
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        if y < 2*cm:
                            c.showPage()
                            y = height - 2*cm
                        
                        c.drawString(2*cm, y, para.text[:80])  # Vereinfacht
                        y -= 0.5*cm
                
                c.save()
                return ConversionResult(True, output_path, ".docx", ".pdf")
            
            else:
                return ConversionResult(False, output_path, ".docx", output_ext,
                                       f"Konvertierung nicht unterstützt: .docx -> {output_ext}")
                
        except Exception as e:
            return ConversionResult(False, output_path, ".docx", output_ext, str(e))
    
    def _convert_from_text(self, input_path: str, output_path: str,
                           output_ext: str) -> ConversionResult:
        """Konvertiert von TXT/MD."""
        input_ext = Path(input_path).suffix.lower()
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(input_path, "r", encoding="latin-1") as f:
                text = f.read()
        
        try:
            if output_ext in {".txt", ".md"}:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
                return ConversionResult(True, output_path, input_ext, output_ext)

            elif output_ext == ".pdf":
                if not REPORTLAB_AVAILABLE:
                    return ConversionResult(False, output_path, input_ext, ".pdf",
                                           "reportlab nicht verfügbar")
                
                c = canvas.Canvas(output_path, pagesize=A4)
                width, height = A4
                y = height - 2*cm
                
                for line in text.split("\n"):
                    if y < 2*cm:
                        c.showPage()
                        y = height - 2*cm
                    
                    c.drawString(2*cm, y, line[:100])
                    y -= 0.4*cm
                
                c.save()
                return ConversionResult(True, output_path, input_ext, ".pdf")
            
            elif output_ext == ".docx":
                if not DOCX_AVAILABLE:
                    return ConversionResult(False, output_path, input_ext, ".docx",
                                           "python-docx nicht verfügbar")
                
                doc = DocxDocument()
                for para in text.split("\n\n"):
                    doc.add_paragraph(para)
                doc.save(output_path)
                return ConversionResult(True, output_path, input_ext, ".docx")
            
            elif output_ext == ".html":
                html = f'<html><head><meta charset="utf-8"></head><body><pre>{text}</pre></body></html>'
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html)
                return ConversionResult(True, output_path, input_ext, ".html")
            
            else:
                return ConversionResult(False, output_path, input_ext, output_ext,
                                       f"Konvertierung nicht unterstützt: {input_ext} -> {output_ext}")
                
        except Exception as e:
            return ConversionResult(False, output_path, input_ext, output_ext, str(e))

    def _convert_from_html(self, input_path: str, output_path: str,
                           output_ext: str) -> ConversionResult:
        """Konvertiert von HTML."""
        import re
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(input_path, "r", encoding="latin-1") as f:
                html_content = f.read()
        
        # Einfache HTML-Tag-Bereinigung für Text/Markdown
        clean_text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        clean_text = re.sub(r'<style[^>]*>.*?</style>', '', clean_text, flags=re.DOTALL | re.IGNORECASE)
        clean_text = re.sub(r'<br\s*/?>', '\n', clean_text, flags=re.IGNORECASE)
        clean_text = re.sub(r'</p>', '\n\n', clean_text, flags=re.IGNORECASE)
        clean_text = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'# \1\n\n', clean_text, flags=re.IGNORECASE)
        clean_text = re.sub(r'<[^>]+>', '', clean_text)
        clean_text = re.sub(r'\n{3,}', '\n\n', clean_text).strip()
        
        try:
            if output_ext in {".txt", ".md"}:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(clean_text)
                return ConversionResult(True, output_path, ".html", output_ext)
            elif output_ext == ".docx":
                if not DOCX_AVAILABLE:
                    return ConversionResult(False, output_path, ".html", ".docx",
                                           "python-docx nicht verfügbar")
                doc = DocxDocument()
                for p in clean_text.split("\n\n"):
                    if p.strip():
                        doc.add_paragraph(p.strip())
                doc.save(output_path)
                return ConversionResult(True, output_path, ".html", ".docx")
            elif output_ext == ".pdf":
                if not REPORTLAB_AVAILABLE:
                    return ConversionResult(False, output_path, ".html", ".pdf",
                                           "reportlab nicht verfügbar")
                c = canvas.Canvas(output_path, pagesize=A4)
                width, height = A4
                y = height - 2*cm
                for line in clean_text.split("\n"):
                    if y < 2*cm:
                        c.showPage()
                        y = height - 2*cm
                    c.drawString(2*cm, y, line[:100])
                    y -= 0.4*cm
                c.save()
                return ConversionResult(True, output_path, ".html", ".pdf")
            else:
                return ConversionResult(False, output_path, ".html", output_ext,
                                       f"Konvertierung nicht unterstützt: .html -> {output_ext}")
        except Exception as e:
            return ConversionResult(False, output_path, ".html", output_ext, str(e))
    
    def _convert_from_image(self, input_path: str, output_path: str,
                            output_ext: str) -> ConversionResult:
        """Konvertiert von Bild."""
        if not PIL_AVAILABLE:
            return ConversionResult(False, output_path, Path(input_path).suffix, output_ext,
                                   "Pillow nicht verfügbar")
        
        input_ext = Path(input_path).suffix.lower()
        try:
            with Image.open(input_path) as img:
                if output_ext == ".pdf":
                    # Bild zu PDF
                    if img.mode in ('RGBA', 'LA', 'PA') or (img.mode == 'P' and 'transparency' in img.info):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode in ('RGBA', 'LA', 'PA'):
                            background.paste(img, mask=img.split()[-1])
                        else:
                            rgba_img = img.convert('RGBA')
                            background.paste(rgba_img, mask=rgba_img.split()[3])
                        img_to_save = background
                    elif img.mode != 'RGB':
                        img_to_save = img.convert('RGB')
                    else:
                        img_to_save = img

                    img_to_save.save(output_path, "PDF")
                    return ConversionResult(True, output_path, input_ext, ".pdf")
                
                elif output_ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}:
                    # Bildformat-Konvertierung
                    if output_ext in {".jpg", ".jpeg", ".bmp"}:
                        if img.mode in ('RGBA', 'LA', 'PA') or (img.mode == 'P' and 'transparency' in img.info):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            if img.mode in ('RGBA', 'LA', 'PA'):
                                background.paste(img, mask=img.split()[-1])
                            else:
                                rgba_img = img.convert('RGBA')
                                background.paste(rgba_img, mask=rgba_img.split()[3])
                            img_to_save = background
                        elif img.mode not in ('RGB', 'L'):
                            img_to_save = img.convert('RGB')
                        else:
                            img_to_save = img
                    else:
                        img_to_save = img
                    
                    format_map = {
                        ".jpg": "JPEG", ".jpeg": "JPEG",
                        ".png": "PNG", ".gif": "GIF",
                        ".bmp": "BMP", ".webp": "WEBP"
                    }
                    img_to_save.save(output_path, format_map.get(output_ext, "PNG"))
                    return ConversionResult(True, output_path, input_ext, output_ext)
                
                else:
                    return ConversionResult(False, output_path, input_ext, output_ext,
                                           f"Konvertierung nicht unterstützt: {input_ext} -> {output_ext}")
                
        except Exception as e:
            return ConversionResult(False, output_path, input_ext, 
                                   output_ext, str(e))


# === Hilfsfunktionen ===

def convert_to_pdf(input_path: str, output_path: Optional[str] = None) -> str:
    """Konvertiert beliebige Datei zu PDF."""
    if output_path is None:
        output_path = str(Path(input_path).with_suffix(".pdf"))
    
    converter = FormatConverter()
    result = converter.convert(input_path, output_path, OutputFormat.PDF)
    
    return result.output_path if result.success else ""


def convert_to_text(input_path: str, output_path: Optional[str] = None) -> str:
    """Konvertiert beliebige Datei zu Text."""
    if output_path is None:
        output_path = str(Path(input_path).with_suffix(".txt"))
    
    converter = FormatConverter()
    result = converter.convert(input_path, output_path, OutputFormat.TXT)
    
    return result.output_path if result.success else ""


def batch_convert(input_paths: List[str], output_dir: str,
                  output_format: OutputFormat) -> List[ConversionResult]:
    """Konvertiert mehrere Dateien."""
    converter = FormatConverter()
    results = []
    
    output_dir_path = Path(output_dir)
    output_dir_path.mkdir(parents=True, exist_ok=True)
    
    for input_path in input_paths:
        output_name = Path(input_path).stem + f".{output_format.value}"
        output_path = str(output_dir_path / output_name)
        
        result = converter.convert(input_path, output_path, output_format)
        results.append(result)
    
    return results
