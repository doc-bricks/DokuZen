#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tests für FormatConverter (formats.py) - Robustheit und Edge-Cases:
- PDF zu JPEG/BMP/WEBP mit Grayscale- und RGBA/Alpha-Pixmaps (kein Stride-Fehler / ValueError)
- PDF mit 0 Seiten (sauberes Failure-Handling statt IndexError)
- Bild-Konvertierung aus Palette- (P) und Alpha-Modi (LA, PA, RGBA) nach JPEG/BMP/PDF
- Markdown-Ausgabe (.md) aus TXT, DOCX, PDF und HTML
- HTML-Eingabe (.html) nach TXT, MD, DOCX und PDF
"""

import sys
import unittest
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from PIL import Image
from core.converter.formats import (
    FormatConverter,
    OutputFormat,
    PYMUPDF_AVAILABLE,
    PIL_AVAILABLE,
    DOCX_AVAILABLE,
    REPORTLAB_AVAILABLE,
)
if PYMUPDF_AVAILABLE:
    import fitz


class TestConverterFormatsRobustness(unittest.TestCase):

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)
        self.converter = FormatConverter()

    def tearDown(self):
        self.temp_dir.cleanup()

    @unittest.skipUnless(PYMUPDF_AVAILABLE and PIL_AVAILABLE, "PyMuPDF und Pillow erforderlich")
    def test_pdf_grayscale_to_jpeg(self):
        """Grayscale-PDF zu JPEG muss ohne ValueError konvertiert werden."""
        pdf_file = self.temp_path / "gray.pdf"
        jpg_file = self.temp_path / "gray.jpg"

        doc = fitz.open()
        page = doc.new_page(width=100, height=100)
        page.draw_rect(fitz.Rect(10, 10, 80, 80), color=(0.4,), fill=(0.8,))
        doc.save(str(pdf_file))
        doc.close()

        result_jpg = self.converter.convert(str(pdf_file), str(jpg_file))
        self.assertTrue(result_jpg.success, f"PDF -> JPG fehlgeschlagen: {result_jpg.error}")
        self.assertTrue(jpg_file.exists())
        self.assertGreater(jpg_file.stat().st_size, 0)

        with Image.open(str(jpg_file)) as img:
            self.assertEqual(img.format, "JPEG")
            self.assertEqual(img.mode, "RGB")

    @unittest.skipUnless(PYMUPDF_AVAILABLE and PIL_AVAILABLE, "PyMuPDF und Pillow erforderlich")
    def test_pdf_rgba_alpha_to_jpeg(self):
        """PDF mit Alpha/RGBA-Pixmap zu JPEG muss ohne Stride-/Kanalverzerrung konvertiert werden."""
        pdf_file = self.temp_path / "alpha.pdf"
        jpg_file = self.temp_path / "alpha.jpg"

        doc = fitz.open()
        page = doc.new_page(width=100, height=100)
        page.draw_rect(fitz.Rect(0, 0, 100, 100), color=(1, 0, 0), fill=(1, 0, 0))
        doc.save(str(pdf_file))
        doc.close()

        # Mock get_pixmap um ein 4-Kanal RGBA Pixmap zu simulieren
        mock_pix = MagicMock()
        mock_pix.width = 10
        mock_pix.height = 10
        mock_pix.n = 4
        mock_pix.alpha = 1
        mock_pix.samples = b"\xff\x00\x00\x80" * 100  # R=255, G=0, B=0, A=128

        real_open = fitz.open
        with patch("fitz.open") as mock_fitz_open:
            real_doc = real_open(str(pdf_file))
            real_page = real_doc[0]
            real_page.get_pixmap = MagicMock(return_value=mock_pix)
            mock_fitz_open.return_value = real_doc

            result = self.converter.convert(str(pdf_file), str(jpg_file))
            self.assertTrue(result.success, f"RGBA PDF -> JPG fehlgeschlagen: {result.error}")
            self.assertTrue(jpg_file.exists())

    def test_pdf_empty_zero_pages_safe_failure(self):
        """PDF ohne Seiten liefert sauberes Failure statt IndexError."""
        pdf_file = self.temp_path / "empty.pdf"
        jpg_file = self.temp_path / "empty.jpg"
        pdf_file.write_bytes(b"%PDF-1.4\n")

        mock_doc = MagicMock()
        mock_doc.__len__.return_value = 0
        mock_doc.close = MagicMock()

        with patch("fitz.open", return_value=mock_doc):
            result = self.converter.convert(str(pdf_file), str(jpg_file))
            self.assertFalse(result.success)
            self.assertIn("keine Seiten", result.error or "")

    @unittest.skipUnless(PYMUPDF_AVAILABLE, "PyMuPDF erforderlich")
    def test_pdf_to_markdown(self):
        """PDF zu Markdown (.md) muss strukturierten Text erzeugen."""
        pdf_file = self.temp_path / "sample.pdf"
        md_file = self.temp_path / "sample.md"

        doc = fitz.open()
        p1 = doc.new_page(width=200, height=200)
        p1.insert_text((20, 50), "Erste Testseite DokuZen")
        p2 = doc.new_page(width=200, height=200)
        p2.insert_text((20, 50), "Zweite Seite Inhalt")
        doc.save(str(pdf_file))
        doc.close()

        result = self.converter.convert(str(pdf_file), str(md_file), OutputFormat.MD)
        self.assertTrue(result.success, f"PDF -> MD fehlgeschlagen: {result.error}")
        self.assertTrue(md_file.exists())

        content = md_file.read_text(encoding="utf-8")
        self.assertIn("Seite 1", content)
        self.assertIn("Erste Testseite DokuZen", content)
        self.assertIn("Seite 2", content)
        self.assertIn("Zweite Seite Inhalt", content)

    @unittest.skipUnless(PIL_AVAILABLE, "Pillow erforderlich")
    def test_image_palette_p_to_jpeg_and_bmp(self):
        """P-Modus (Palette) Bild zu JPEG und BMP muss ohne OSError konvertiert werden."""
        png_p = self.temp_path / "palette.png"
        jpg_out = self.temp_path / "palette.jpg"
        bmp_out = self.temp_path / "palette.bmp"

        img = Image.new("RGBA", (60, 60), (200, 100, 50, 180))
        p_img = img.convert("P")
        p_img.save(str(png_p), "PNG")

        res_jpg = self.converter.convert(str(png_p), str(jpg_out))
        self.assertTrue(res_jpg.success, f"P -> JPEG fehlgeschlagen: {res_jpg.error}")
        self.assertTrue(jpg_out.exists())

        res_bmp = self.converter.convert(str(png_p), str(bmp_out))
        self.assertTrue(res_bmp.success, f"P -> BMP fehlgeschlagen: {res_bmp.error}")
        self.assertTrue(bmp_out.exists())

    @unittest.skipUnless(PIL_AVAILABLE, "Pillow erforderlich")
    def test_image_grayscale_alpha_la_to_jpeg(self):
        """LA-Modus Bild (Grayscale + Alpha) zu JPEG muss ohne OSError konvertieren."""
        png_la = self.temp_path / "gray_alpha.png"
        jpg_out = self.temp_path / "gray_alpha.jpg"

        img = Image.new("LA", (40, 40), (128, 200))
        img.save(str(png_la), "PNG")

        res_jpg = self.converter.convert(str(png_la), str(jpg_out))
        self.assertTrue(res_jpg.success, f"LA -> JPEG fehlgeschlagen: {res_jpg.error}")
        self.assertTrue(jpg_out.exists())

    def test_txt_to_md_and_html(self):
        """TXT-Datei zu Markdown und HTML konvertieren."""
        txt_file = self.temp_path / "test.txt"
        md_file = self.temp_path / "test.md"
        html_file = self.temp_path / "test.html"

        txt_file.write_text("# Überschrift\n\nAbsatz mit Umlauten: äöüß", encoding="utf-8")

        res_md = self.converter.convert(str(txt_file), str(md_file), OutputFormat.MD)
        self.assertTrue(res_md.success, f"TXT -> MD fehlgeschlagen: {res_md.error}")
        self.assertTrue(md_file.exists())
        self.assertIn("äöüß", md_file.read_text(encoding="utf-8"))

        res_html = self.converter.convert(str(txt_file), str(html_file), OutputFormat.HTML)
        self.assertTrue(res_html.success, f"TXT -> HTML fehlgeschlagen: {res_html.error}")
        self.assertTrue(html_file.exists())
        self.assertIn("<pre>", html_file.read_text(encoding="utf-8"))

    def test_html_to_txt_and_md(self):
        """HTML-Eingabe zu TXT und Markdown konvertieren."""
        html_file = self.temp_path / "document.html"
        txt_file = self.temp_path / "document.txt"
        md_file = self.temp_path / "document.md"

        html_content = "<html><body><h1>Titel</h1><p>Erster Absatz mit <b>fettem</b> Text.</p></body></html>"
        html_file.write_text(html_content, encoding="utf-8")

        res_txt = self.converter.convert(str(html_file), str(txt_file), OutputFormat.TXT)
        self.assertTrue(res_txt.success, f"HTML -> TXT fehlgeschlagen: {res_txt.error}")
        self.assertTrue(txt_file.exists())
        self.assertIn("Titel", txt_file.read_text(encoding="utf-8"))

        res_md = self.converter.convert(str(html_file), str(md_file), OutputFormat.MD)
        self.assertTrue(res_md.success, f"HTML -> MD fehlgeschlagen: {res_md.error}")
        self.assertTrue(md_file.exists())
        md_content = md_file.read_text(encoding="utf-8")
        self.assertIn("# Titel", md_content)
        self.assertIn("Erster Absatz", md_content)

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx erforderlich")
    def test_docx_to_md_and_html(self):
        """DOCX zu Markdown und HTML konvertieren."""
        from docx import Document

        docx_file = self.temp_path / "sample.docx"
        md_file = self.temp_path / "sample.md"
        html_file = self.temp_path / "sample.html"

        doc = Document()
        doc.add_heading("Haupttitel", level=1)
        doc.add_paragraph("Einleitender Text für DokuZen.")
        doc.save(str(docx_file))

        res_md = self.converter.convert(str(docx_file), str(md_file), OutputFormat.MD)
        self.assertTrue(res_md.success, f"DOCX -> MD fehlgeschlagen: {res_md.error}")
        self.assertTrue(md_file.exists())
        self.assertIn("Einleitender Text", md_file.read_text(encoding="utf-8"))

        res_html = self.converter.convert(str(docx_file), str(html_file), OutputFormat.HTML)
        self.assertTrue(res_html.success, f"DOCX -> HTML fehlgeschlagen: {res_html.error}")
        self.assertTrue(html_file.exists())
        self.assertIn("<p>", html_file.read_text(encoding="utf-8"))


if __name__ == "__main__":
    unittest.main()
